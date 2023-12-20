from fastapi import FastAPI, Body, Response, Cookie
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from starlette.responses import RedirectResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles



import random
from random import randint

import docx
from docx.shared import Cm
from docx2pdf import convert

from openpyxl import load_workbook
from openpyxl import Workbook
from datetime  import date

import os

import conf
from conf import mk_DOCX
from conf import mk_XL
from conf import DB
from conf import DataBase
from conf import Auth

from pypdf import PdfMerger

import json

from datetime import datetime

from typing import Optional



def make_clean(Id):
    pul = os.listdir(path='.')
    if len(pul) >= 10:
        for file in pul:
            if "Tula" in file:
                os.remove(file)
    for file in pul:
        if f"Tula{Id}" in file:
            os.remove(file)



def get_para_data(output_doc_name, paragraph):
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
   


app = FastAPI()

app.mount("/static", StaticFiles(directory="public", html=True))
app.mount("/static", StaticFiles(directory="/", html=True))



origins = [
    "https://217.144.103.121",
    "https://localhost:8080",
    "https://emk.websto.pro"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST"],#, "OPTIONS", "DELETE", "PATH", "PUT"],
    allow_headers=["Content-Type", "Accept", "Location", "Allow", "Content-Disposition", "Sec-Fetch-Dest"],
)



@app.get("/")
def root():
    return RedirectResponse("https://emk.websto.pro/static/select.html")

@app.get("/ep4")
def ep4():
    return RedirectResponse("https://emk.websto.pro/static/ep4.html")

@app.get("/epn")
def epn():
    return RedirectResponse("https://emk.websto.pro/static/epn.html")

@app.get("/ep4Reductor")
def ep4Reductor():
    return RedirectResponse("https://emk.websto.pro/static/ep4Reductor.html")

@app.get("/classic")
def ep4Reductor():
    return RedirectResponse("https://emk.websto.pro/static/ep4Reductor.html")

@app.get("/vimu")
def ep4Reductor():
    return RedirectResponse("https://emk.websto.pro/static/vimu.html")



@app.get("/Tula/{ID}/{name}/{mark}", response_class = FileResponse)
def download_file(ID, name, mark, token: Optional[str] = Cookie(default=None)):
    #headers = {'Content-Disposition': f'attachment; filename="Tula{ID}.pdf"'}
    auth = Auth()
    ans = auth.get_user(token)
    if ans["valid"]:
        os.system(f"cp Tula{ID}.docx /files")
        auth.auth_file(ID, f'ОЛ ТЭП {name} {ID} {mark}.docx', ans["user"]["user_id"])
    return FileResponse(f'Tula{ID}.docx', filename=f'ОЛ ТЭП {name} {ID} {mark}.docx', media_type='application/docx')#, headers=headers)

@app.get("/user")
def user_data(token: Optional[str] = Cookie(default=None)):
    auth = Auth()
    ans = auth.get_user(token)

    print(ans)
    
    if ans["valid"]:
        print("Токен валидный")
    else:
        print("Токен не валидный")
    
    return ans


@app.get("/TulaEXEL/{ID}/{name}/{mark}", response_class = FileResponse)
def download_file(ID, name, mark, token: Optional[str] = Cookie(default=None)):

    #читаем куки и проверяем валидность токена
    auth = Auth()
    ans = auth.get_user(token)
    if ans["valid"]:
        os.system(f"cp Tula{ID}.xlsx /files")
        auth.auth_file(ID, f'ТКП {name} {ID} {mark}.xlsx', ans["user"]["user_id"])

        return FileResponse(f'Tula{ID}.xlsx', filename=f'ТКП {name} {ID} {mark}.xlsx', media_type='application/xlsx')#, headers=headers)
    else:
        return RedirectResponse("https://emk.websto.pro/static/select.html")

@app.get("/TulaPDF/{ID}/{name}/{mark}", response_class = FileResponse)
def download_file(ID, name, mark, token: Optional[str] = Cookie(default=None)):
    os.system(f"unoconv -f pdf Tula{ID}.xlsx")
    os.rename(f"Tula{ID}.pdf", "Tula.pdf")
    os.system(f"unoconv -f pdf Tula{ID}.docx")

    pdfs = ["Tula.pdf", f"Tula{ID}.pdf"]
    merger = PdfMerger()
    for pdf in pdfs:
         merger.append(pdf)
    merger.write(f"Tula{ID}.pdf")
    merger.close()

    #читаем куки и проверяем валидность токена
    auth = Auth()
    ans = auth.get_user(token)

    if ans["valid"]:
        os.system(f"cp Tula{ID}.pdf /files")
        auth.auth_file(ID, f'Информация на подпись  {name} {ID} {mark}.pdf', ans["user"]["user_id"])

        return FileResponse(f'Tula{ID}.pdf', filename=f'Информация на подпись  {name} {ID} {mark}.pdf', media_type='application/pdf')
    else:
        return RedirectResponse("https://emk.websto.pro/static/select.html")

@app.get("/LK")
def ep4Reductor():
    return RedirectResponse("https://emk.websto.pro/static/lk.html")



@app.get("/user_files/{ID}")
def get_file(ID):
    auth = Auth()
    ans = auth.get_files(ID)

    return ans

@app.get("/getFile/{ID}")
def download_file(ID, token: Optional[str] = Cookie(default=None)):
    auth=Auth()
    file_name = auth.download_file(ID)
    fl_nm = file_name[0]
    if fl_nm[-3:] == "pdf":
        fl_rormat = "pdf"
    else:
        fl_rormat = fl_nm[-4:]
    return FileResponse(f'/files/{fl_nm}', filename=file_name[1], media_type=f'application/{fl_rormat}')

@app.post("/DBep")
def get_param(jsn = Body()):
    a = jsn["a"]

    for i in range(9-len(a)):
        a.append("")

    if a[0] == 'ЭПН':
        a.insert(2, '')
    print(a)

    db = DataBase()
    res = db.get_params(mark = a[0], Isp = a[1], flc_type = a[2], Mom = a[3], N = a[4], V = a[5], sh = a[6], flc = a[7], fz = a[8])
    print(res)
    return {"ans" : res}

@app.post("/DBRN")
def get_param(jsn = Body()):
    a = jsn["a"]
    for i in range(8-len(a)):
        a.append("")
    db = DataBase()
    res=db.get_RN(rn = a[0], isp = a[1], flc = a[2], Mom = a[3], V = a[4], t = a[5], sh = a[6])
    print(res)
    return {"ans" : res}

@app.post("/DBclassic")
def get_param(jsn = Body()):
    a = jsn["a"]
    for i in range(10-len(a)):
        a.append("")
    db = DataBase()
    res=db.get_classic(isp = a[0], flc = a[1], N = a[2], ob = a[3], V = a[4], Mom = a[5], bkw = a[6], elpod = a[7], nom = a[8])
    print(res)
    return {"ans" : res}

@app.post("/m1")
def get_param(jsn = Body()):
    a = jsn["a"]
    db = DataBase()
    ans = db.get_M1(a[0], a[1], a[2], a[3])
    
    return {"ans" : ans}



@app.post("/DB")
def get_param(jsn = Body()):
    a = jsn["a"]
    bd = DB(a[0])
    for i in range(8-len(a)):
        a.append("")
    if a[0] == "ЭП4":
        res = bd.get_params(a[0], a[1], a[2], a[3], a[4], a[5], "", "")
    elif a[0] == "ЭПН":
        res = bd.get_params(a[0], a[1], a[2], a[3], a[4], a[5], a[6], a[7])
    try:
        res.sort(key=float)
    except:
        pass
    print(res)
    return {"ans" : res}

@app.post("/DBrn")
def get_param(jsn = Body()):
    a = jsn["a"]
    bd = DB("ЭП4")
    for i in range(7-len(a)):
        a.append("")
    ans = bd.get_RN(a[0], a[1], a[2], a[3], a[4], "")
    res=[]
    for x in ans:
        if (x != " ") and (x != "null") and (x != None):
            res.append(x)
    try:
        res.sort(key=float)
    except:
        pass
    print(res)
    return {"ans" : res}

@app.post("/classicDB")
def get_param(jsn = Body()):
    a = jsn["a"]
    bd = DB("ЭП4")
    for i in range(8-len(a)):
        a.append("")
    ans = bd.get_classic(a[0], a[1], a[2], a[3], a[4], a[5], a[6], a[7])
    res=[]
    for x in ans:
        if (x != " ") and (x != "null") and (x != None):
            res.append(x)
    try:
        res.sort(key=float)
    except:
        pass

    print(res)
    return {"ans" : res}

@app.post("/M1")
def get_param(jsn = Body()):
    a = jsn["a"]
    bd = DB("M1")
    ans = bd.get_M1(a[0], a[1], a[2], a[3])
    
    return {"ans" : ans}

@app.post("/Mark")
def get_mark(jsn = Body()):
    a = jsn["a"]
    for i in range(8-len(a)):
        a.append("")
    print(a)
    bd = DataBase()
    ans = bd.getMark(a)
    print(ans)
    return {"ans" : ans}

@app.post("/login")
def login(jsn = Body()):
    user = jsn
    print(user)
    #запрос на БД
    auth = Auth()
    tkn = auth.get_token(user)
    #записать в куки
    return {"token" : tkn}

@app.get("/pdf/{file}")
def auth(file: str):
    return RedirectResponse(f"https://emk.websto.pro/pdf/{file}")

@app.get("/pdf/conSchemes/{file}")
def auth(file: str):
    return RedirectResponse(f"https://emk.websto.pro/pdf/conSchemes/{file}")

@app.get("/{token}")
def auth(token: str, response: Response):
    auth = Auth()
    ans = auth.get_user(token)

    response = RedirectResponse(f"https://emk.websto.pro/")
    if ans["valid"]:
        #вписать токен в куки
        response.set_cookie(key="token", value=token, samesite=None)

        #response.headers["token"] = token
        print("Токен валидный")
    else:
        print("Токен не валидный")
    return response

@app.post("/download")
def download_file(jsn = Body()):

    '''Чтение и форматирование JSON'''
    jsn0 = jsn["jsn0"]
    jsn1 = jsn["jsn1"]
    jsn2 = jsn["jsn2"]
    jsn3 = jsn["jsn3"]
    jsn4 = jsn["jsn4"]
    jsn5 = jsn["jsn5"]
    jsn6 = jsn["jsn6"]
    jsn7 = jsn["jsn7"]

    import json
    f = open("ID.json", 'r')
    ID = json.load(f)
    f.close()

    ID = int(ID["ID"]) + 1
    
    
    
    mark = jsn1[1]
    print(jsn2[1])

    if mark[:3] == "ЭП4":
        dc = mk_DOCX(jsn5[1], jsn0, jsn1, jsn2, jsn3, jsn4, jsn6) 
        tbls = dc.ep4()

        xl = mk_XL(ID, mark, [jsn0, jsn1, jsn2, jsn3, jsn4, jsn5, jsn6, jsn7])
        wb = xl.ep()

    elif mark[:3] == "ЭПН":
        dc = mk_DOCX(jsn5[1], jsn0, jsn1, jsn2, jsn3, jsn4, jsn6) 
        tbls = dc.epn()

        xl = mk_XL(ID, mark, [jsn0, jsn1, jsn2, jsn3, jsn4, jsn5, jsn6, jsn7])
        wb = xl.ep()

    elif mark[:4] == 'ВИМУ':
        dc = mk_DOCX(jsn5[1], jsn0, jsn1, jsn2, jsn3, jsn4, [jsn5[1], jsn5[3]])
        tbls = dc.vimu()

        xl = mk_XL(ID, mark, [jsn0, jsn2, jsn3, [jsn4[0], jsn4[2], jsn4[4], jsn4[5], jsn5[1], jsn5[3]] ])
        wb = xl.vimu()

    else:
        dc = mk_DOCX(jsn5[1], jsn0, jsn1, jsn2, jsn3, jsn4,  jsn5)
        tbls = dc.classic()#jsn6[0])

        xl = mk_XL(ID, mark, [jsn0, jsn1, jsn2, jsn3, jsn4, jsn5, jsn6])
        wb = xl.classic()

    '''Работа с .docx'''
    doc = docx.Document('Опросный_лист_ТЭП.docx')
    doc_new = docx.Document()
    
    doc_new.add_picture('img.png', width=Cm(1.76), height=Cm(1.67))
    if mark[:4] != "ВИМУ":
        for para in doc.paragraphs:
            if para.text == 'ОРГАНИЗАЦИЯ - заказчик: ______________________________________________________':
                para.text = f'ОРГАНИЗАЦИЯ - заказчик: \t {tbls["ans3"][0]}'
            if para.text == 'Ф.И.О. ____________________ Тел: ____________________ E-mail:_____________________':
                para.text = f' Ф.И.О. \t  {tbls["ans3"][1]} \n Тел.: \t  {tbls["ans3"][2]} \n E-mail:  {tbls["ans3"][3]} \n Дата: «   » ____________________ 202__ г'
            
            get_para_data(doc_new, para)
            if para.text == 'Характеристика и параметры арматуры:':
                tbl1 = doc_new.add_table(rows=0, cols=2)
                tbl1.style = 'Table Grid'
                ks=list(tbls["names1"])
                vs=list(tbls["ans1"])
                for i in range(len(ks)):
                    row_cells = tbl1.add_row().cells
                    row_cells[0].text = ks[i]
                    cell = tbl1.cell(i, 1)
                    cell.text = vs[i]
                #doc_new.add_paragraph()
                
            if para.text == 'Характеристика и параметры электропривода:':
                tbl1 = doc_new.add_table(rows=0, cols=2)
                tbl1.style = 'Table Grid'
                ks=list(tbls["names2"])
                vs=list(tbls["ans2"])
                for i in range(len(ks)):
                    row_cells = tbl1.add_row().cells
                    row_cells[0].text = ks[i]
                    cell = tbl1.cell(i, 1)
                    cell.text = vs[i]
    
    elif mark[:4] == 'ВИМУ':
        for para in doc.paragraphs:
            if para.text == 'ОРГАНИЗАЦИЯ - заказчик: ______________________________________________________':
                txt = f'ОРГАНИЗАЦИЯ - заказчик: \t {tbls["ans3"][0]}'
                p = doc_new.add_paragraph('')
                run = p.add_run(txt)

            elif para.text == 'Ф.И.О. ____________________ Тел: ____________________ E-mail:_____________________':
                txt = f' Ф.И.О. \t  {tbls["ans3"][1]} \n Тел.: \t  {tbls["ans3"][2]} \n E-mail:  {tbls["ans3"][3]} \n Дата: «   » ____________________ 202__ г'
                p = doc_new.add_paragraph('')
                run = p.add_run(txt)

            elif para.text == 'Характеристика и параметры электропривода:':
                para.text = ""

            elif para.text == 'Характеристика и параметры арматуры:':
                txt = 'Характеристика и параметры внешнего интеллектуального модуля управления'
                p = doc_new.add_paragraph('')
                run = p.add_run(txt)
                
                tbl1 = doc_new.add_table(rows=0, cols=2)
                tbl1.style = 'Table Grid'

                ks=list(tbls["names2"])
                vs=list(tbls["ans2"])
                for i in range(len(ks)):
                    row_cells = tbl1.add_row().cells
                    row_cells[0].text = ks[i]
                    cell = tbl1.cell(i, 1)
                    cell.text = vs[i]
                tbl1 = doc_new.add_paragraph('')
            
            else:
                get_para_data(doc_new, para)

    
    
    '''Работа с файлами'''
    #сборщик мусора
    make_clean(ID)

    #сохранение файлов с уникальным id
    doc_new.save(f'Tula{ID}.docx')
    wb.save(f"Tula{ID}.xlsx")

    jsn = {"ID" : ID} 

    with open("ID.json", "w") as fh:
        json.dump(jsn, fh)
    #convert(f"Tula{ID}.docx")

    file_break=['/', '\\', '\n', '*', '|', '\'', '\"']
    Mark = ""
    k=0
    for m in mark:
        if m == "-":
            k+=1
        if (k < 4) and (m not in file_break):
            Mark += m
    print(Mark)
        


    return {"id" : ID, "name" : jsn0[1], "mark" : Mark}