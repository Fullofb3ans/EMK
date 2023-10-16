from fastapi import FastAPI, Body
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from starlette.responses import RedirectResponse, Response
from fastapi.staticfiles import StaticFiles



import random
from random import randint

import docx
from docx.shared import Cm
from docx2pdf import convert

from openpyxl import load_workbook
from openpyxl import Workbook
from datetime  import date

import os

import conf
from conf import mk_DOCX
from conf import mk_XL
from conf import DB



def make_clean(Id):
    pul = os.listdir(path='.')
    if len(pul) >= 10:
        for file in pul:
            if "Tula" in file:
                os.remove(file)
    for file in pul:
        if f"Tula{Id}" in file:
            os.remove(file)



def get_para_data(output_doc_name, paragraph):
    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
   


app = FastAPI()

app.mount("/static", StaticFiles(directory="public", html=True))
app.mount("/static", StaticFiles(directory="/", html=True))



origins = [
    "https://217.144.103.121",
    "https://localhost:8080",
    "https://emk.websto.pro"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST"],#, "OPTIONS", "DELETE", "PATH", "PUT"],
    allow_headers=["Content-Type", "Accept", "Location", "Allow", "Content-Disposition", "Sec-Fetch-Dest"],
)



@app.get("/")
def root():
    return RedirectResponse("https://emk.websto.pro/static/select.html")

@app.get("/ep4")
def ep4():
    return RedirectResponse("https://emk.websto.pro/static/ep4.html")

@app.get("/epn")
def epn():
    return RedirectResponse("https://emk.websto.pro/static/epn.html")

@app.get("/ep4Reductor")
def ep4Reductor():
    return RedirectResponse("https://emk.websto.pro/static/ep4Reductor.html")



@app.get("/Tula/{ID}", response_class = FileResponse)
def download_file(ID):
    #headers = {'Content-Disposition': f'attachment; filename="Tula{ID}.pdf"'}
    return FileResponse(f'Tula{ID}.docx', filename=f'ОЛ ТЭП {ID}.docx', media_type='application/docx')#, headers=headers)

@app.get("/TulaEXEL/{ID}", response_class = FileResponse)
def download_file(ID):
    return FileResponse(f'Tula{ID}.xlsx', filename=f'ТКП {ID}.xlsx', media_type='application/xlsx')#, headers=headers)



@app.post("/DB")
def get_param(jsn = Body()):
    a = jsn["a"]
    bd = DB(a[0])
    for i in range(7-len(a)):
        a.append("")
    if a[0] == "ЭП4":
        res = bd.get_params(a[0], a[1], a[2], a[3], a[4], "", "")
    elif a[0] == "ЭПН":
        res = bd.get_params(a[0], a[1], a[2], a[3], a[4], a[5], a[6])
    try:
        res.sort(key=int)
    except:
        res = [res]
    print(res)
    return {"ans" : res}



@app.post("/DBrn")
def get_param(jsn = Body()):
    a = jsn["a"]
    bd = DB("ЭП4")
    for i in range(6-len(a)):
        a.append("")
    ans = bd.get_RN(a[0], a[1], a[2], a[3], a[4], "")
    res=[]
    for x in ans:
        if (x != " ") and (x != "null") and (x != None):
            res.append(x)
    try:
        res.sort(key=int)
    except:
        pass
    print(res)
    return {"ans" : res}



@app.post("/download")
def download_file(jsn = Body()):

    '''Чтение и форматирование JSON'''
    jsn0 = jsn["jsn0"]
    jsn1 = jsn["jsn1"]
    jsn2 = jsn["jsn2"]
    jsn3 = jsn["jsn3"]
    jsn4 = jsn["jsn4"]
    jsn5 = jsn["jsn5"]
    jsn6 = jsn["jsn6"]
    jsn7 = jsn["jsn7"]

    ID = randint(1000000, 10000000)
    
    
    
    mark = jsn1[1]

    if mark[:3] == "ЭП4":
        dc = mk_DOCX(jsn0, jsn1, jsn2, jsn3, jsn4, jsn6) 
        tbls = dc.ep4()

        xl = mk_XL(ID, [jsn0, jsn1, jsn2, jsn3, jsn4, jsn5, jsn6, jsn7])
        wb = xl.ep()
    elif mark[:3] == "ЭПН":
        dc = mk_DOCX(jsn0, jsn1, jsn2, jsn3, jsn4, jsn6) 
        tbls = dc.epn()

        xl = mk_XL(ID, [jsn0, jsn1, jsn2, jsn3, jsn4, jsn5, jsn6, jsn7])
        wb = xl.ep()
    elif mark[:4] == 'ВИМУ':
        dc = mk_DOCX(jsn0, jsn1, jsn2, jsn3, jsn4, [jsn5[1], jsn5[3]])
        tbls = dc.vimu()

        xl = mk_XL(ID, [jsn0, jsn2, jsn3, [jsn4[0], jsn4[2], jsn4[4], jsn4[5], jsn5[1], jsn5[3]] ])
        wb = xl.vimu()



    '''Работа с .docx'''
    doc = docx.Document('Опросный_лист_ТЭП.docx')
    doc_new = docx.Document()
    
    doc_new.add_picture('img.png', width=Cm(1.76), height=Cm(1.67))

    for para in doc.paragraphs:
        if para.text == 'ОРГАНИЗАЦИЯ - заказчик: ______________________________________________________':
            para.text = f'ОРГАНИЗАЦИЯ - заказчик: \t {tbls["ans3"][0]}'
        if para.text == 'Ф.И.О. ____________________ Тел: ____________________ E-mail:_____________________':
            para.text = f'Ф.И.О. \t  {tbls["ans3"][1]} \n Тел.: \t  {tbls["ans3"][2]} \n E-mail:  {tbls["ans3"][3]} \n \t\t\t \t\t\t \t\t Дата: «   » ____________________ 202__ г'
        
        get_para_data(doc_new, para)
        if mark[:4] != 'ВИМУ':
            if para.text == 'Характеристика и параметры арматуры:':
                tbl1 = doc_new.add_table(rows=0, cols=2)
                tbl1.style = 'Table Grid'
                ks=list(tbls["names1"])
                vs=list(tbls["ans1"])
                for i in range(len(ks)):
                    row_cells = tbl1.add_row().cells
                    row_cells[0].text = ks[i]
                    cell = tbl1.cell(i, 1)
                    cell.text = vs[i]
                doc_new.add_paragraph()
                
            if para.text == 'Характеристика и параметры электропривода:':
                tbl1 = doc_new.add_table(rows=0, cols=2)
                tbl1.style = 'Table Grid'
                ks=list(tbls["names2"])
                vs=list(tbls["ans2"])
                for i in range(len(ks)):
                    row_cells = tbl1.add_row().cells
                    row_cells[0].text = ks[i]
                    cell = tbl1.cell(i, 1)
                    cell.text = vs[i]

        elif mark[:4] == 'ВИМУ':
            if para.text == 'Характеристика и параметры арматуры:':
                para.text = 'Характеристика и параметры внешнего интеллектуального модуля управления'

                tbl1 = doc_new.add_table(rows=0, cols=2)
                tbl1.style = 'Table Grid'


                ks=list(tbls["names2"])
                vs=list(tbls["ans2"])
                for i in range(len(ks)):
                    row_cells = tbl1.add_row().cells
                    row_cells[0].text = ks[i]
                    cell = tbl1.cell(i, 1)
                    cell.text = vs[i]
            
            if para.text == 'Характеристика и параметры электропривода:':
                para.text = ""

    
    '''Работа с файлами'''
    #сборщик мусора
    make_clean(ID)

    #сохранение файлов с уникальным id
    doc_new.save(f'Tula{ID}.docx')
    wb.save(f"Tula{ID}.xlsx")
    #convert(f"Tula{ID}.docx")



    return {"id" : ID}